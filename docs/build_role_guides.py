from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

import build_user_guide as guide_base
from build_user_guide import (
    AMBER,
    INK,
    LINE,
    MINT,
    MUTED,
    PALE_AMBER,
    PALE_RED,
    RED,
    TEAL,
    TEAL_DARK,
    WHITE,
    add_bullets,
    add_page_number,
    add_simple_table,
    add_steps,
    add_text,
    make_styles,
    rgb,
    set_cell_border,
    set_cell_shading,
    set_font,
    set_table_geometry,
)


# TH Sarabun New is installed on the deployment workstation and preserves Thai
# glyphs in both Microsoft Word and the LibreOffice-based QA renderer.
guide_base.FONT = "TH Sarabun New"

_base_set_font = set_font


def set_font(run, size=None, bold=None, color=None, italic=None):
    """Apply the Thai-capable typeface and mark the run language explicitly."""
    _base_set_font(run, size=size, bold=bold, color=color, italic=italic)
    lang = run._element.get_or_add_rPr().find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.get_or_add_rPr().append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:eastAsia"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")


# Helper functions imported from the shared guide builder resolve set_font from
# their defining module, so point that shared reference at the localized wrapper.
guide_base.set_font = set_font


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
LOGO = ROOT / "logo.png"
TEACHER_OUTPUT = DOCS / "คู่มือการใช้งาน CRS MIS — ครูผู้สอน.docx"
ADMIN_OUTPUT = DOCS / "คู่มือการใช้งาน CRS MIS — ฝ่ายวิชาการ (Admin).docx"


# compact_reference_guide preset with two named project overrides:
# 1) TH Sarabun New replaces Calibri to preserve Thai in Word and rendered QA.
# 2) CRS MIS teal replaces preset blue for established product branding.
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def configure_compact_lists(doc):
    """Resolve compact-reference list geometry into styles + numbering XML."""
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        p_style = abstract.find(".//" + qn("w:pStyle"))
        if p_style is None or p_style.get(qn("w:val")) not in {"ListBullet", "ListNumber"}:
            continue
        level = abstract.find(qn("w:lvl"))
        if level is None:
            continue
        p_pr = level.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            level.append(p_pr)
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        for child in list(tabs):
            tabs.remove(child)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "270")
        tabs.append(tab)
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")


def add_callout(doc, title, body, kind="info"):
    palette = {
        "info": (MINT, TEAL_DARK),
        "caution": (PALE_AMBER, AMBER),
        "danger": (PALE_RED, RED),
    }
    fill, title_color = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=TABLE_INDENT_DXA)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "18", "color": title_color},
        top={"val": "single", "sz": "4", "color": fill},
        bottom={"val": "single", "sz": "4", "color": fill},
        end={"val": "single", "sz": "4", "color": fill},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title), size=11, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    set_font(p2.add_run(body), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_role_cover(doc, role, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(3.1))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("คู่มือการปฏิบัติงาน"), size=11, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("CRS MIS"), size=28, bold=True, color=TEAL_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_font(p.add_run(role), size=20, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_font(p.add_run("ระบบสารสนเทศเพื่อการบริหารโรงเรียนเจริญศึกษา"), size=13, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [7200], indent_dxa=1080)
    cell = table.cell(0, 0)
    set_cell_shading(cell, MINT)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "4", "color": LINE},
        bottom={"val": "single", "sz": "4", "color": LINE},
        start={"val": "single", "sz": "4", "color": LINE},
        end={"val": "single", "sz": "4", "color": LINE},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(description), size=11, color=INK)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run("ฉบับปรับปรุง 20 กรกฎาคม 2569"), size=9.5, color=MUTED)
    doc.add_page_break()


def setup_document(title, subject, running_label):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "CRS MIS"
    props.comments = "คู่มือการใช้งานระบบสารสนเทศเพื่อการบริหารโรงเรียนเจริญศึกษา"

    make_styles(doc)
    for style in doc.styles:
        if style.type != 1:  # WD_STYLE_TYPE.PARAGRAPH
            continue
        r_pr = style._element.get_or_add_rPr()
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:val"), "th-TH")
        lang.set(qn("w:eastAsia"), "th-TH")
        lang.set(qn("w:bidi"), "th-TH")
    configure_compact_lists(doc)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run(running_label), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    add_page_number(footer)
    return doc


def add_teacher_overview(doc):
    doc.add_heading("1. เริ่มต้นใช้งาน", level=1)
    add_text(doc, "คู่มือนี้สำหรับครูผู้สอนและครูผู้รับผิดชอบกิจกรรม ใช้เป็นลำดับงานตั้งแต่เช็คชื่อ กรอกคะแนน ไปจนถึงตรวจผลและติดตามนักเรียน")
    add_simple_table(
        doc,
        ["งาน", "เมนูที่ใช้"],
        [
            ("เช็คชื่อและชั่วโมงชดเชย", "เช็คชื่อ"),
            ("กรอกคะแนนและติด/ถอด ร.", "กรอกคะแนน"),
            ("ตรวจคะแนน เกรด ร. และ มส.", "สรุปผลการเรียน"),
            ("ติดตามการขาดเรียน", "เฝ้าระวัง มส."),
            ("กำหนด/ให้คะแนนกิจกรรมที่รับผิดชอบ", "กำหนดสมรรถนะ / ลงคะแนนสมรรถนะ"),
        ],
        [3200, 6160],
    )
    doc.add_heading("1.1 เข้าสู่ระบบและเลือกปี", level=2)
    add_steps(doc, [
        "เปิดเว็บไซต์ CRS MIS และเข้าสู่ระบบด้วยอีเมล/รหัสผ่านที่ฝ่ายวิชาการออกให้",
        "เข้าฝ่ายวิชาการ ระบบจะเริ่มที่หน้าเช็คชื่อ",
        "ตรวจปีการศึกษาที่เลือกทุกครั้ง โดยระบบจะจำปีเดียวกันไว้ใช้ข้ามหน้า",
        "เลือกชั้น ห้อง และวิชาตามลำดับก่อนโหลดรายชื่อ",
    ])
    add_callout(doc, "สิทธิ์ของครู", "ครูแก้ไขได้เฉพาะวิชาหรือรายการสมรรถนะที่ตนเป็นเจ้าของ หากเปิดรายการของผู้อื่น ระบบจะแสดงแบบอ่านอย่างเดียวหรือไม่โหลดข้อมูลสำหรับแก้ไข", "caution")
    doc.add_heading("1.2 ลำดับงานแนะนำ", level=2)
    add_bullets(doc, [
        "ทุกวัน: เช็คชื่อทันทีหลังเริ่มคาบ และตรวจจำนวนคาบให้ถูกต้อง",
        "หลังประเมินงาน: กรอกคะแนนตามหน่วย/องค์ประกอบ แล้วเปิดหน้าสรุปเพื่อตรวจยอด",
        "ทุกสัปดาห์: เปิดเฝ้าระวัง มส. ดูนักเรียนที่ขาดสะสมตั้งแต่ระดับเสี่ยง",
        "ก่อนส่งผล: ตรวจ ร. มส. คะแนนรวม และความครบของคะแนนสมรรถนะ",
    ])


def add_teacher_attendance(doc):
    doc.add_page_break()
    doc.add_heading("2. เช็คชื่อและชั่วโมงชดเชย", level=1)
    doc.add_heading("2.1 เช็คชื่อ", level=2)
    add_steps(doc, [
        "เลือกปีการศึกษา ชั้น ห้อง และวิชาพื้นฐาน",
        "เลือกวันที่และจำนวนคาบของครั้งนั้น แล้วกดโหลดรายชื่อ",
        "กำหนดสถานะรายคน: มา ขาด ลาป่วย ลากิจ หรือสาย",
        "ตรวจชื่อและจำนวนคาบอีกครั้ง แล้วกดบันทึก",
        "ถ้าต้องแก้ครั้งเดิม ให้เลือกวิชาและวันที่เดิม ระบบจะโหลดข้อมูลมาให้แก้ไข",
    ])
    add_simple_table(
        doc,
        ["สถานะ", "ผลต่อยอดขาด"],
        [
            ("มา / สาย", "ไม่นับเป็นขาด"),
            ("ขาด", "นับเต็มตามจำนวนคาบ"),
            ("ลาป่วย / ลากิจ", "นับครึ่งหนึ่งของจำนวนคาบ"),
            ("ชั่วโมงชดเชย", "นำไปลดจำนวนคาบขาดสุทธิเมื่อเข้าเงื่อนไข"),
        ],
        [2400, 6960],
    )
    add_callout(doc, "วิชาบูรณาการ", "หน้าเช็คชื่อให้เลือกวิชาพื้นฐานสมาชิกที่ครูรับผิดชอบ การเช็คชื่อจะถูกเก็บแยกตามวิชาพื้นฐานจริง", "info")
    doc.add_heading("2.2 บันทึกชั่วโมงชดเชย", level=2)
    add_steps(doc, [
        "ในหน้าเช็คชื่อ เลื่อนมาที่ส่วนชั่วโมงชดเชย",
        "เลือกนักเรียนและวิชาที่ทำงาน/เรียนเพิ่มเติมจริง",
        "กรอกจำนวนคาบพร้อมบันทึก",
        "เปิดสรุปผลการเรียนเพื่อตรวจว่าจำนวนคาบเข้าเรียนสุทธิได้รับการปรับแล้ว",
    ])
    add_callout(doc, "ข้อจำกัด", "หากขาดเกิน 40% ของคาบทั้งเทอม ระบบถือเป็นกรณีเรียนซ้ำรายวิชาและไม่ใช้ชั่วโมงชดเชยกู้ผล มส.", "danger")


def add_teacher_scores(doc):
    doc.add_page_break()
    doc.add_heading("3. กรอกคะแนนและจัดการ ร.", level=1)
    doc.add_heading("3.1 กรอกคะแนน", level=2)
    add_steps(doc, [
        "เปิดกรอกคะแนน แล้วเลือกปี ชั้น ห้อง และวิชาพื้นฐาน",
        "เลือกประเภทคะแนน: คะแนนวิชา หรือสมรรถนะหลัก",
        "เลือกหน่วยใหญ่ หน่วยย่อย และครั้งประเมินตามลำดับ",
        "กรอกคะแนนดิบรายคนโดยไม่เกินคะแนนเต็มของครั้งนั้น",
        "บันทึก แล้วเปิดสรุปผลการเรียนเพื่อตรวจผลที่ระบบเทียบสัดส่วนให้อัตโนมัติ",
    ])
    add_callout(
        doc,
        "หน้าที่ของครูเจ้าของวิชา",
        "ครูเจ้าของวิชาต้องกรอกคะแนนสมรรถนะหลักของนักเรียนในวิชาที่รับผิดชอบด้วยตนเอง โดยเลือกประเภทคะแนน “สมรรถนะหลัก” และกรอกให้ครบตามโครงสร้างที่กำหนด การกรอกเฉพาะคะแนนวิชายังถือว่างานไม่ครบ",
        "caution",
    )
    add_bullets(doc, [
        "รายชื่อมีเฉพาะนักเรียนที่ลงทะเบียนในวิชานั้น",
        "วิชาบูรณาการไม่กรอกคะแนนตรง ระบบคำนวณจากวิชาพื้นฐานสมาชิก",
        "สมรรถนะหลักแสดงแยก 6 ด้าน ด้านละเต็ม 100 ไม่รวมเป็นเลขเดียว",
    ])
    doc.add_heading("3.2 ติดหรือถอด ร.", level=2)
    add_steps(doc, [
        "เลือกวิชาในหน้ากรอกคะแนน แล้วไปที่การ์ด “ร. (ไม่จบหลักสูตร)”",
        "เลือกนักเรียน กรอกเหตุผลให้ชัดเจน แล้วกดติด ร.",
        "เมื่อส่งงานหรือหลักฐานครบ ให้เลือกรายการเดิมและกดถอด ร.",
        "ตรวจผลอีกครั้งในหน้าสรุปผลการเรียน",
    ])
    add_callout(doc, "ใช้ ร. ให้ถูกกรณี", "ร. ใช้เมื่อผลงานหรือหลักฐานการเรียนรู้ยังไม่ครบ ไม่ใช้แทนปัญหาการเข้าเรียน ซึ่งต้องพิจารณาผ่านเกณฑ์ มส.", "caution")


def add_teacher_reports(doc):
    doc.add_page_break()
    doc.add_heading("4. ตรวจผลและติดตามนักเรียน", level=1)
    doc.add_heading("4.1 สรุปผลการเรียน", level=2)
    add_steps(doc, [
        "เลือกปี ชั้น ห้อง และวิชา",
        "เลือกดูรายคนหรือทั้งห้อง",
        "ตรวจคะแนนหน่วย คะแนนรวม สมรรถนะ การเข้าเรียน และผลเกรด/ร./มส.",
        "หากเป็นวิชาบูรณาการ ให้ตรวจผลรวมจากวิชาพื้นฐานสมาชิก",
    ])
    doc.add_heading("4.2 เฝ้าระวัง มส.", level=2)
    add_simple_table(
        doc,
        ["ระดับ", "เกณฑ์เตือน", "สิ่งที่ควรทำ"],
        [
            ("เสี่ยง", "ขาดสะสมตั้งแต่ 10%", "ติดต่อและติดตามการมาเรียน"),
            ("วิกฤต / มส.", "ขาดสะสมเกิน 20%", "วางแผนเรียนเพิ่มเติมหรือชดเชย"),
            ("เรียนซ้ำรายวิชา", "ขาดสะสมเกิน 40%", "ประสานฝ่ายวิชาการดำเนินการ"),
        ],
        [1700, 2800, 4860],
    )
    add_callout(doc, "การคำนวณช่วงต้นเทอม", "ระบบเทียบจำนวนคาบขาดสะสมจริงกับเพดานของคาบทั้งเทอม จึงไม่ตัดสิน มส. จากเปอร์เซ็นต์ของคาบที่เช็คไปเพียงไม่กี่ครั้ง", "info")
    doc.add_heading("4.3 รายงานเรียนซ้ำชั้น", level=2)
    add_text(doc, "ครูสามารถเปิดรายงานเรียนซ้ำชั้นเพื่อดูธงเตือนเบื้องต้น แต่การยืนยันผลและการดำเนินการเป็นหน้าที่ของฝ่ายวิชาการ")


def add_teacher_competency(doc):
    doc.add_page_break()
    doc.add_heading("5. สมรรถนะจากกิจกรรมและกิจวัตร", level=1)
    doc.add_heading("5.1 กำหนดรายการที่รับผิดชอบ", level=2)
    add_steps(doc, [
        "เปิดกำหนดสมรรถนะ สร้างรายการ ระบุปี ภาคเรียน ประเภท และชื่อ",
        "ใน “รายการที่กำหนดไว้” แตะหัวกลุ่มกิจกรรมหรือกิจวัตรเพื่อกางรายการ",
        "เลือกรายการ แล้วเพิ่มสมรรถนะ องค์ประกอบมาตรฐาน และคะแนนเต็ม",
        "เพิ่มผู้เข้าร่วมทั้งห้องหรือเลือกเป็นรายคน",
    ])
    add_bullets(doc, [
        "ครูจัดการได้เฉพาะรายการที่ตนรับผิดชอบ ส่วน Admin จัดการได้ทุกรายการ",
        "หากมีครั้งประเมินแล้ว ปี ประเภท และภาคเรียนจะถูกล็อกเพื่อรักษาประวัติ",
        "รายชื่อผู้เข้าร่วมอิงชั้น/ห้องของปีรายการ ไม่ใช่ชั้นปัจจุบันหลังเลื่อนชั้น",
    ])
    doc.add_heading("5.2 ลงคะแนน", level=2)
    add_steps(doc, [
        "เปิดลงคะแนนสมรรถนะ เลือกปี ภาคเรียน และรายการ",
        "เลือกช่วงชั้นและองค์ประกอบ แล้วสร้างหรือเลือกครั้งประเมิน",
        "เปิดคะแนนรายนักเรียน กรอกคะแนนดิบ และบันทึก",
        "ตรวจความครบของคะแนนก่อนจบภาคเรียน",
    ])
    add_callout(doc, "Snapshot รายชื่อ", "เมื่อสร้างครั้งประเมิน ระบบจะเก็บ snapshot ผู้เข้าร่วมไว้ ประวัติเดิมจึงไม่เปลี่ยนแม้นักเรียนเลื่อนชั้นหรือมีการแก้สมาชิกภายหลัง", "info")


def add_teacher_support(doc):
    doc.add_page_break()
    doc.add_heading("6. ตรวจงานและแก้ปัญหา", level=1)
    doc.add_heading("6.1 เช็กลิสต์ก่อนออกจากระบบ", level=2)
    add_bullets(doc, [
        "ปี ชั้น ห้อง และวิชาที่เลือกถูกต้อง",
        "จำนวนคาบและสถานะการเข้าเรียนครบทุกคน",
        "คะแนนไม่เกินคะแนนเต็มและกดบันทึกแล้ว",
        "ตรวจ ร. และนักเรียนเสี่ยง มส.",
        "ออกจากระบบเมื่อใช้เครื่องร่วมกับผู้อื่น",
    ])
    doc.add_heading("6.2 ปัญหาที่พบบ่อย", level=2)
    add_simple_table(
        doc,
        ["อาการ", "ตรวจสอบ"],
        [
            ("ไม่พบรายชื่อนักเรียน", "ตรวจปี/ชั้น/ห้อง/วิชา และแจ้ง Admin ตรวจการลงทะเบียน"),
            ("แก้ไขไม่ได้", "ตรวจว่าเป็นเจ้าของวิชาหรือรายการนั้น"),
            ("ไม่พบวิชา", "ตรวจปี ภาคเรียน ชั้น และลองตัวเลือกทุกปีสำหรับข้อมูลเก่า"),
            ("บันทึกชดเชยไม่ได้", "ตรวจว่ายอดขาดไม่เกินเกณฑ์เรียนซ้ำ 40%"),
            ("ลืมรหัสผ่าน", "ติดต่อฝ่ายวิชาการให้รีเซ็ตผ่าน Supabase Dashboard"),
        ],
        [3000, 6360],
    )
    add_callout(doc, "แจ้งปัญหาให้ตรวจเร็ว", "ส่งชื่อหน้า ปีการศึกษา วิชา ชื่อนักเรียน (ถ้ามี) และภาพหน้าจอให้ฝ่ายวิชาการ โดยไม่ส่งรหัสผ่าน", "info")


def build_teacher_guide():
    doc = setup_document(
        "คู่มือการใช้งาน CRS MIS — ครูผู้สอน",
        "คู่มือปฏิบัติงานสำหรับครูผู้สอนและครูผู้รับผิดชอบกิจกรรม",
        "CRS MIS | คู่มือครูผู้สอน",
    )
    add_role_cover(doc, "สำหรับครูผู้สอน", "เช็คชื่อ · กรอกคะแนน · ติดตามผล · ประเมินสมรรถนะ")
    add_teacher_overview(doc)
    add_teacher_attendance(doc)
    add_teacher_scores(doc)
    add_teacher_reports(doc)
    add_teacher_competency(doc)
    add_teacher_support(doc)
    doc.save(TEACHER_OUTPUT)
    return TEACHER_OUTPUT


def add_admin_overview(doc):
    doc.add_heading("1. บทบาทและหลักความปลอดภัย", level=1)
    add_text(doc, "คู่มือนี้สำหรับฝ่ายวิชาการที่มีสิทธิ์ Admin ครอบคลุมการเตรียมข้อมูล ควบคุมสิทธิ์ ตรวจคุณภาพข้อมูล และขึ้นปีการศึกษาใหม่")
    add_simple_table(
        doc,
        ["ขอบเขต", "ความรับผิดชอบหลัก"],
        [
            ("บัญชีและสิทธิ์", "สร้างบัญชีครู กำหนดเจ้าของวิชา/รายการ และจำกัด Admin เท่าที่จำเป็น"),
            ("ข้อมูลหลัก", "นักเรียน ประวัติชั้น/ห้อง วิชา โครงสร้างคะแนน และการลงทะเบียน"),
            ("กำกับผล", "ตรวจรายงาน ร. มส. เรียนซ้ำชั้น และความครบของสมรรถนะ"),
            ("เปลี่ยนปี", "เลื่อนชั้น จบ/ย้ายออก ยกวิชา และตรวจข้อมูลปีใหม่"),
        ],
        [2500, 6860],
    )
    add_callout(doc, "หลักสำคัญ", "ข้อมูลส่วนใหญ่แก้ย้อนหลังได้ แต่การลบนักเรียนและการขึ้นปีใหม่กระทบข้อมูลจำนวนมาก ควรตรวจ preview สำรองข้อมูล และยืนยันขอบเขตก่อนดำเนินการ", "danger")
    doc.add_heading("1.1 เริ่มต้นระบบ", level=2)
    add_steps(doc, [
        "สร้างบัญชีครูใน Supabase Dashboard และกำหนด role ใน profiles",
        "ตั้งปีการศึกษาที่ใช้งาน และตรวจชั้นสูงสุดของโรงเรียน",
        "เพิ่มนักเรียนพร้อมปี ชั้น และห้อง",
        "สร้างวิชา กำหนดครูเจ้าของ และลงทะเบียนนักเรียน",
        "ทดสอบด้วยบัญชีครูหนึ่งบัญชีก่อนเปิดใช้งานทั้งโรงเรียน",
    ])


def add_admin_students(doc):
    doc.add_page_break()
    doc.add_heading("2. จัดการนักเรียน", level=1)
    doc.add_heading("2.1 เพิ่มและตรวจรายชื่อ", level=2)
    add_steps(doc, [
        "เปิดนักเรียน ระบุปีการศึกษาของชั้นและห้อง",
        "เพิ่มทีละคน หรือวางจาก Excel 5 คอลัมน์: ชั้น ห้อง รหัส ชื่อ นามสกุล",
        "ตรวจ preview และแก้แถวผิดก่อนบันทึก",
        "ใช้ตัวกรองชั้น/ห้อง แล้วแตะหัวกลุ่ม เช่น ม.3/1 เพื่อกางรายชื่อ",
    ])
    add_bullets(doc, [
        "รายชื่อแบ่งเป็นกลุ่มชั้น/ห้องและพับเริ่มต้น เพื่อลดความยาวเมื่อมีนักเรียนทั้งโรงเรียน",
        "ติ๊กแสดงนักเรียนที่จบ/ย้ายออกเมื่อต้องตรวจประวัติ",
        "students.grade_level/classroom คือสถานะปัจจุบัน ส่วนประวัติรายปีอยู่ใน student_year_placements",
    ])
    doc.add_heading("2.2 ลบนักเรียน", level=2)
    add_steps(doc, [
        "กางห้องและกดไอคอนลบของนักเรียนที่ต้องการ",
        "อ่านรายการข้อมูลที่ผูกอยู่ เช่น คะแนน เช็คชื่อ ร. ชดเชย สมรรถนะ และการลงทะเบียน",
        "หากมีข้อมูลผูกอยู่ ให้พิมพ์ชื่อนักเรียนให้ตรงเพื่อยืนยัน",
    ])
    add_callout(doc, "ลบแล้วกู้คืนไม่ได้", "การลบนักเรียนจะ cascade ลบข้อมูลที่เกี่ยวข้องทั้งหมด หากเพียงไม่ต้องการให้ใช้งานต่อ ให้ใช้สถานะจบหรือย้ายออกตอนขึ้นปีใหม่แทน", "danger")


def add_admin_subjects(doc):
    doc.add_page_break()
    doc.add_heading("3. วิชา โครงสร้างคะแนน และการลงทะเบียน", level=1)
    doc.add_heading("3.1 สร้างวิชาพื้นฐาน", level=2)
    add_steps(doc, [
        "เปิดจัดการโครงสร้าง เลือกปี ชั้น และภาคเรียน",
        "สร้างวิชา ระบุรหัส ชื่อ ประเภท ระดับชั้น คะแนนเต็ม คาบทั้งเทอม และครูเจ้าของ",
        "เพิ่มหน่วยใหญ่ หน่วยย่อย และครั้งประเมิน",
        "เพิ่มนักเรียนทั้งห้องหรือรายคนเข้าสู่ enrollments",
        "ทดสอบหน้ากรอกคะแนนและเช็คชื่อว่ามีรายชื่อครบ",
    ])
    doc.add_heading("3.2 วิชาบูรณาการ", level=2)
    add_bullets(doc, [
        "สร้างเป็นประเภทบูรณาการและเลือกวิชาพื้นฐานสมาชิกจากปีเดียวกัน",
        "ไม่สร้างโครงสร้างคะแนนหรือกรอกคะแนนตรงในวิชาบูรณาการ",
        "ผลคะแนนใช้การถ่วงน้ำหนักจากวิชาสมาชิก และเวลาเรียนรวมจากวิชาพื้นฐาน",
        "ห้ามลบวิชาพื้นฐานที่ยังเป็นสมาชิกของวิชาบูรณาการ",
    ])
    add_callout(doc, "ก่อนเปิดให้ครูใช้", "ทุกวิชาต้องมีปี ครูเจ้าของ จำนวนคาบ และรายชื่อลงทะเบียนครบ มิฉะนั้นหน้าปฏิบัติงานจะว่างหรือคำนวณ มส. ไม่ได้", "caution")


def add_admin_competency(doc):
    doc.add_page_break()
    doc.add_heading("4. กำกับสมรรถนะ", level=1)
    doc.add_heading("4.1 ค่ากลาง", level=2)
    add_bullets(doc, [
        "กำหนดน้ำหนักรายวิชา กิจกรรม และกิจวัตรแยกแต่ละสมรรถนะ โดยรวมเป็น 100%",
        "กำหนดช่วงคะแนนและชื่อระดับการแปลผลให้ต่อเนื่องตั้งแต่ 1–100",
        "ตรวจองค์ประกอบมาตรฐานตามช่วงชั้นก่อนให้ครูสร้างรายการ",
    ])
    doc.add_heading("4.2 รายการกิจกรรม/กิจวัตร", level=2)
    add_steps(doc, [
        "กรองปีการศึกษาในส่วนรายการที่กำหนดไว้",
        "แตะหัวกลุ่มกิจกรรมหรือกิจวัตรเพื่อกางรายการ",
        "ตรวจชื่อ ปี ภาคเรียน ครูผู้รับผิดชอบ องค์ประกอบ และผู้เข้าร่วม",
        "หากต้องเปลี่ยนผู้รับผิดชอบ ให้แก้ไขรายการด้วยสิทธิ์ Admin",
    ])
    add_callout(doc, "รายการที่มีคะแนนแล้ว", "เมื่อมีครั้งประเมิน ระบบล็อกปี ประเภท และภาคเรียน การลบจะลบองค์ประกอบ สมาชิก snapshot และคะแนนทั้งหมด จึงต้องตรวจผลกระทบก่อนยืนยัน", "danger")


def add_admin_reports(doc):
    doc.add_page_break()
    doc.add_heading("5. รายงานและการควบคุมคุณภาพ", level=1)
    add_simple_table(
        doc,
        ["รายงาน", "ใช้ตรวจ"],
        [
            ("สรุปผลการเรียน", "คะแนนรวม เกรด ร. มส. เวลาเรียน และวิชาบูรณาการ"),
            ("เฝ้าระวัง มส.", "นักเรียนเสี่ยง/วิกฤตจากทุกวิชาในห้อง"),
            ("เรียนซ้ำชั้น", "GPA ต่ำกว่า 1.00 หรือวิชา 0/ร./มส. มากกว่าครึ่ง"),
            ("สรุปสมรรถนะ", "ความครบของคะแนนจากแหล่งที่มีน้ำหนักมากกว่า 0%"),
        ],
        [2800, 6560],
    )
    doc.add_heading("5.1 รอบตรวจแนะนำ", level=2)
    add_bullets(doc, [
        "รายสัปดาห์: ตรวจเฝ้าระวัง มส. และประสานครูที่ปรึกษา",
        "กลางภาค: ตรวจคะแนนที่ยังว่าง ร. และจำนวนคาบทั้งเทอม",
        "ก่อนส่งผล: ติดตามให้ครูเจ้าของวิชากรอกสมรรถนะหลักของวิชาตนเอง แล้วตรวจเกรด/ร./มส. และความครบของสมรรถนะทุกแหล่ง",
        "ก่อนขึ้นปีใหม่: ตรวจ placement นักเรียน วิชาปีต้นทาง และรายการจบ/ย้ายออก",
    ])


def add_admin_rollover(doc):
    doc.add_page_break()
    doc.add_heading("6. ขึ้นปีการศึกษาใหม่", level=1)
    add_callout(doc, "ขอบเขตการทำงาน", "การขึ้นปีใหม่ดำเนินการกับนักเรียน active ทุกคนพร้อมกัน ไม่ใช่เฉพาะกลุ่มที่กำลังกางดู รายการที่พับยังอยู่ใน DOM และจะถูกประมวลผลตามค่าที่เลือก", "danger")
    doc.add_heading("6.1 ก่อนเริ่ม", level=2)
    add_bullets(doc, [
        "สำรองฐานข้อมูลหรือ snapshot สถานะนักเรียน",
        "ยืนยันชั้นสูงสุดที่เปิดสอน",
        "ตรวจปีต้นทางและปีใหม่ว่าไม่ซ้ำ",
        "ตรวจนักเรียนซ้ำชั้น ย้ายออก และจบการศึกษากับผู้รับผิดชอบ",
        "ตรวจรายชื่อวิชาที่ควรยกและวิชาที่จะสร้างโครงสร้างใหม่",
    ])
    doc.add_heading("6.2 ขั้นตอน", level=2)
    add_steps(doc, [
        "เลือกปีต้นทางและกรอกปีใหม่",
        "ในเลื่อนชั้นนักเรียน แตะหัวชั้นเพื่อกาง ตรวจสรุปสด และตั้งค่ารายคน/ทั้งชั้น",
        "ในยกวิชาไปปีใหม่ แตะหัวชั้นเพื่อกาง เลือกหรือไม่เลือกวิชารายตัว/ทั้งชั้น",
        "ตรวจสรุปรวมจำนวนเลื่อน จบ ซ้ำชั้น ย้ายออก และจำนวนวิชาที่จะยก",
        "กดตรวจสอบและเริ่มขึ้นปีใหม่เพียงครั้งเดียว แล้วยืนยันข้อความเตือน",
        "รอจน log แสดงว่าเสร็จ จากนั้นตรวจปีใหม่ นักเรียน วิชา placement และ enrollment",
    ])
    add_callout(doc, "ห้ามปิดกลางคัน", "กระบวนการก๊อปด้วย JavaScript ไม่เป็น transaction เดียว หากขัดข้องให้เก็บ log และตรวจข้อมูลที่สร้างไปแล้วก่อนแก้ไข ห้ามกดรันซ้ำทันที", "caution")


def add_admin_support(doc):
    doc.add_page_break()
    doc.add_heading("7. แก้ปัญหาและเช็กลิสต์ Admin", level=1)
    add_simple_table(
        doc,
        ["อาการ", "แนวทางตรวจ"],
        [
            ("ครูแก้วิชาไม่ได้", "ตรวจ owner_id และ role ของบัญชี"),
            ("ไม่มีรายชื่อในวิชา", "ตรวจ enrollments และปี/ชั้น/ห้องของวิชา"),
            ("ห้องหรือรายงานย้อนหลังผิด", "ตรวจ student_year_placements ของปีที่เลือก"),
            ("คำนวณ มส. ผิด", "ตรวจ total_periods, attendance และ makeup_hours"),
            ("สรุปสมรรถนะไม่ออก", "ตรวจน้ำหนักและคะแนนทุกแหล่งที่มีน้ำหนักมากกว่า 0"),
            ("ขึ้นปีใหม่รันซ้ำไม่ได้", "ตรวจวิชาปีใหม่ที่สร้างค้างและ log ก่อนดำเนินการต่อ"),
        ],
        [3000, 6360],
    )
    doc.add_heading("7.1 เช็กลิสต์ก่อนเปิดใช้งาน", level=2)
    add_bullets(doc, [
        "บัญชีครูและสิทธิ์ถูกต้อง",
        "ปีการศึกษาที่ใช้งานถูกต้องทุกหน้า",
        "นักเรียนมี placement และสถานะ active ถูกต้อง",
        "วิชามีเจ้าของ คาบเรียน โครงสร้าง และ enrollment",
        "ครูทดลองเช็คชื่อ/กรอกคะแนนหนึ่งรายการและเห็นผลในสรุป",
        "มีช่องทางรับแจ้งปัญหาและผู้รับผิดชอบสำรอง",
    ])
    add_callout(doc, "ข้อมูลที่ควรขอเมื่อรับแจ้ง", "ชื่อหน้า บัญชีผู้ใช้ ปี ชั้น ห้อง วิชา/รายการ ชื่อนักเรียน เวลาเกิดปัญหา ข้อความแจ้งเตือน และภาพหน้าจอ โดยห้ามขอรหัสผ่านจากครู", "info")


def build_admin_guide():
    doc = setup_document(
        "คู่มือการใช้งาน CRS MIS — ฝ่ายวิชาการ (Admin)",
        "คู่มือปฏิบัติงานสำหรับผู้ดูแลระบบฝ่ายวิชาการ",
        "CRS MIS | คู่มือฝ่ายวิชาการ (Admin)",
    )
    add_role_cover(doc, "สำหรับฝ่ายวิชาการ (Admin)", "ตั้งค่าระบบ · จัดการข้อมูล · ตรวจรายงาน · ขึ้นปีการศึกษาใหม่")
    add_admin_overview(doc)
    add_admin_students(doc)
    add_admin_subjects(doc)
    add_admin_competency(doc)
    add_admin_reports(doc)
    add_admin_rollover(doc)
    add_admin_support(doc)
    doc.save(ADMIN_OUTPUT)
    return ADMIN_OUTPUT


def main():
    teacher = build_teacher_guide()
    admin = build_admin_guide()
    print(teacher)
    print(admin)


if __name__ == "__main__":
    main()
